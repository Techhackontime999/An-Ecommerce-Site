import io
import os

from django.conf import settings
from django.shortcuts import render
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.html import strip_tags
from weasyprint import HTML
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Pt, RGBColor, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import DocumentationSection

SECTION_ICONS = {
    'getting-started': 'fa-compass',
    'seller-guide': 'fa-store',
    'payment-methods': 'fa-credit-card',
}
DEFAULT_SECTION_ICON = 'fa-file-lines'

PROJECT_DIR = getattr(settings, 'PROJECT_DIR', os.path.dirname(settings.BASE_DIR))
FONT_DIR = os.path.join(PROJECT_DIR, 'shop', 'static', 'fonts', 'outfit')

# ─── Brand tokens ────────────────────────────────────────
OUTFIT = 'Outfit'
MONO = 'Consolas'
INK = RGBColor(0x0F, 0x17, 0x2A)
ACCENT = RGBColor(0xF9, 0x73, 0x16)
ACCENT_DARK = RGBColor(0xEA, 0x58, 0x0C)
BODY = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_HEX = 'F97316'
ACCENT_SOFT_HEX = 'FFF7ED'
CODE_BG_HEX = 'F1F5F9'


def _prepare_sections():
    sections = DocumentationSection.objects.all()
    total_words = 0
    for section in sections:
        section.disp_icon = SECTION_ICONS.get(section.slug, DEFAULT_SECTION_ICON)
        total_words += len(strip_tags(section.content or '').split())
    read_time = max(1, round(total_words / 200))
    return sections, read_time


def documentation_view(request):
    sections, read_time = _prepare_sections()
    return render(request, 'documentation/documentation.html', {
        'documentation_sections': sections,
        'read_time': read_time,
    })


def download_pdf(request):
    sections, read_time = _prepare_sections()
    html = render_to_string('documentation/pdf_template.html', {
        'documentation_sections': sections,
        'read_time': read_time,
        'today': timezone.localtime(timezone.now()),
        'font_dir': FONT_DIR,
    })
    pdf = HTML(string=html).write_pdf()

    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="shop-seed-documentation.pdf"'
    return response


def download_word(request):
    sections = list(DocumentationSection.objects.all())
    doc = _build_docx(sections)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = 'attachment; filename="shop-seed-documentation.docx"'
    return response


# ─── Word helpers ────────────────────────────────────────
def _set_run(run, size=10.5, color=BODY, bold=False, italic=False,
             mono=False, name=OUTFIT, small_caps=False, spacing=None):
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    font = MONO if mono else name
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:cs'), font)
    rfonts.set(qn('w:eastAsia'), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if small_caps:
        run.font.small_caps = True
    if spacing:
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(spacing))
        sz = rpr.find(qn('w:sz'))
        if sz is not None:
            sz.addprevious(sp)
        else:
            rpr.append(sp)


def _style_paragraph_font(style, size=10.5, color=BODY, bold=False):
    style.font.name = OUTFIT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), OUTFIT)
    rfonts.set(qn('w:hAnsi'), OUTFIT)
    rfonts.set(qn('w:cs'), OUTFIT)
    rfonts.set(qn('w:eastAsia'), OUTFIT)


def _add_hyperlink(paragraph, url, text, size=10.5):
    if not url:
        run = paragraph.add_run(text)
        _set_run(run, size=size, color=ACCENT_DARK, bold=True)
        return run
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), OUTFIT)
    rfonts.set(qn('w:hAnsi'), OUTFIT)
    rpr.append(rfonts)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'EA580C')
    rpr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.append(underline)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rpr.append(sz)
    run.append(rpr)
    text_el = OxmlElement('w:t')
    text_el.set(qn('xml:space'), 'preserve')
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _style_pbdr(p, edges):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge, attrs in edges.items():
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), attrs.get('val', 'single'))
        el.set(qn('w:sz'), str(attrs.get('sz', 8)))
        el.set(qn('w:space'), str(attrs.get('space', 4)))
        el.set(qn('w:color'), attrs.get('color', ACCENT_HEX))
        pbdr.append(el)
    pstyle = ppr.find(qn('w:pStyle'))
    if pstyle is not None:
        pstyle.addnext(pbdr)
    else:
        ppr.insert(0, pbdr)


def _shade_paragraph(p, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    ppr = p._p.get_or_add_pPr()
    pbdr = ppr.find(qn('w:pBdr'))
    if pbdr is not None:
        pbdr.addnext(shd)
        return
    pstyle = ppr.find(qn('w:pStyle'))
    if pstyle is not None:
        pstyle.addnext(shd)
    else:
        ppr.insert(0, shd)


def _shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcw = tcpr.find(qn('w:tcW'))
    if tcw is not None:
        tcw.addnext(shd)
    else:
        tcpr.insert(0, shd)


def _add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    return run


def _add_rich_inline(paragraph, node, size=10.5, color=BODY, bold=False, italic=False):
    for child in node.children:
        name = getattr(child, 'name', None)
        if name is None:
            text = str(child).replace('\xa0', ' ')
            if text:
                run = paragraph.add_run(text)
                _set_run(run, size=size, color=color, bold=bold, italic=italic)
        elif name == 'br':
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif name in ('strong', 'b'):
            _add_rich_inline(paragraph, child, size=size, color=color, bold=True, italic=italic)
        elif name in ('em', 'i'):
            _add_rich_inline(paragraph, child, size=size, color=color, bold=bold, italic=True)
        elif name == 'code':
            text = child.get_text()
            if text:
                run = paragraph.add_run(text)
                _set_run(run, size=max(7.5, size - 1.5), color=ACCENT_DARK, mono=True)
        elif name == 'a':
            _add_hyperlink(paragraph, child.get('href') or '', child.get_text(), size=size)
        else:
            _add_rich_inline(paragraph, child, size=size, color=color, bold=bold, italic=italic)


def _add_rich_block(doc, el):
    name = el.name
    if name == 'p':
        p = doc.add_paragraph()
        _add_rich_inline(p, el)
    elif name in ('h1', 'h2', 'h3', 'h4'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        sizes = {'h1': 14, 'h2': 13, 'h3': 11.5, 'h4': 10.5}
        _add_rich_inline(p, el, size=sizes.get(name, 12), color=INK, bold=True)
    elif name in ('ul', 'ol'):
        style = 'List Bullet' if name == 'ul' else 'List Number'
        for li in el.find_all('li', recursive=False):
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(3)
            _add_rich_inline(p, li)
    elif name == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        _style_pbdr(p, {'left': {'sz': 18, 'space': 6, 'color': ACCENT_HEX}})
        _shade_paragraph(p, ACCENT_SOFT_HEX)
        _add_rich_inline(p, el, color=INK)
    elif name == 'pre':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        _shade_paragraph(p, CODE_BG_HEX)
        run = p.add_run(el.get_text())
        _set_run(run, size=8.5, color=INK, mono=True)
    elif name == 'table':
        rows = el.find_all('tr')
        if rows:
            cols = max(1, len(rows[0].find_all(['th', 'td'])))
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            for i, tr in enumerate(rows):
                cells = tr.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    if j >= cols:
                        break
                    target = table.rows[i].cells[j]
                    target.paragraphs[0].text = ''
                    _add_rich_inline(target.paragraphs[0], cell)
                    if cell.name == 'th':
                        _shade_cell(target, CODE_BG_HEX)


def _add_richtext(doc, html):
    soup = BeautifulSoup(html or '', 'html.parser')
    for el in soup.children:
        if getattr(el, 'name', None):
            _add_rich_block(doc, el)
        else:
            text = str(el).strip()
            if not text:
                continue
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    _set_run(run)


def _build_docx(sections):
    doc = Document()

    cp = doc.core_properties
    cp.title = 'Shop-Seed Documentation'
    cp.author = 'Shop-Seed'
    cp.subject = 'Platform guide for buyers and sellers'

    # Base styles
    normal = doc.styles['Normal']
    _style_paragraph_font(normal, size=10.5, color=BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for style_name in ('List Bullet', 'List Number'):
        try:
            _style_paragraph_font(doc.styles[style_name], size=10.5, color=BODY)
        except KeyError:
            pass

    heading = doc.styles['Heading 1']
    _style_paragraph_font(heading, size=16, color=INK, bold=True)
    heading.paragraph_format.space_before = Pt(20)
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.keep_with_next = True

    # ─── Cover section ───
    cover = doc.sections[0]
    cover.page_width = Mm(210)
    cover.page_height = Mm(297)
    cover.top_margin = Cm(2.0)
    cover.bottom_margin = Cm(2.0)
    cover.left_margin = Cm(2.0)
    cover.right_margin = Cm(2.0)
    cover.footer.is_linked_to_previous = False

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('SHOP-SEED')
    _set_run(run, size=18, color=INK, bold=True, spacing=80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run('DOCUMENTATION  ·  v1.0')
    _set_run(run, size=9, color=ACCENT_DARK, bold=True, small_caps=True, spacing=40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('Everything you need to get started')
    _set_run(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    p.paragraph_format.left_indent = Cm(2.5)
    p.paragraph_format.right_indent = Cm(2.5)
    run = p.add_run(
        'Step-by-step guides for buyers and sellers — from placing your '
        'first order to growing a full store on Shop-Seed.'
    )
    _set_run(run, size=11, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    _style_pbdr(p, {'bottom': {'sz': 12, 'space': 4, 'color': ACCENT_HEX}})

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('%d sections   ·   ~%d min read   ·   %s' % (
        len(sections),
        max(1, round(sum(len(strip_tags(s.content or '').split()) for s in sections) / 200)),
        timezone.localtime(timezone.now()).strftime('%d %b %Y'),
    ))
    _set_run(run, size=9.5, color=MUTED)

    # ─── Content section (new section → footer with page numbers) ───
    content = doc.add_section(WD_SECTION.NEW_PAGE)
    content.page_width = Mm(210)
    content.page_height = Mm(297)
    content.top_margin = Cm(2.4)
    content.bottom_margin = Cm(2.2)
    content.left_margin = Cm(2.2)
    content.right_margin = Cm(2.2)
    content.header_distance = Cm(1.1)
    content.footer_distance = Cm(1.1)
    content.footer.is_linked_to_previous = False

    footer = content.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(16.8), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run('Shop-Seed  ·  shopseed.com')
    _set_run(r, size=8, color=MUTED)
    fp.add_run('\t')
    r = fp.add_run('Page ')
    _set_run(r, size=8, color=MUTED)
    _add_field(fp, 'PAGE')
    r = fp.add_run(' of ')
    _set_run(r, size=8, color=MUTED)
    _add_field(fp, 'NUMPAGES')

    # ─── Table of contents ───
    toc = doc.add_paragraph(style='Heading 1')
    _style_pbdr(toc, {'left': {'sz': 22, 'space': 6, 'color': ACCENT_HEX}})
    run = toc.add_run('Table of Contents')
    _set_run(run, size=16, color=INK, bold=True)

    for i, sec in enumerate(sections, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run('%02d    ' % i)
        _set_run(run, size=11, color=ACCENT_DARK, bold=True)
        run = p.add_run(sec.title)
        _set_run(run, size=11, color=INK, bold=True)
        p.add_run('\t')
        run = p.add_run('Section %02d' % i)
        _set_run(run, size=9.5, color=MUTED)

    # ─── Sections ───
    for i, sec in enumerate(sections, 1):
        heading_p = doc.add_paragraph(style='Heading 1')
        _style_pbdr(heading_p, {'left': {'sz': 22, 'space': 6, 'color': ACCENT_HEX}})
        run = heading_p.add_run('%02d  ' % i)
        _set_run(run, size=16, color=ACCENT, bold=True)
        run = heading_p.add_run(sec.title)
        _set_run(run, size=16, color=INK, bold=True)

        _add_richtext(doc, sec.content)

        if sec.image and hasattr(sec.image, 'path') and os.path.exists(sec.image.path):
            try:
                doc.add_picture(sec.image.path, width=Cm(15))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last.paragraph_format.space_before = Pt(8)
            except Exception:
                pass

        if i < len(sections):
            divider = doc.add_paragraph()
            divider.paragraph_format.space_after = Pt(12)
            _style_pbdr(divider, {'bottom': {'sz': 4, 'space': 1, 'color': 'E2E8F0'}})

    # ─── Closing ───
    if sections:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(28)
        run = p.add_run('Thank you for choosing Shop-Seed')
        _set_run(run, size=14, color=INK, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('This guide is a living document. Need more help with your store or an order?')
        _set_run(run, size=10, color=BODY)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('support@shopseed.com    ·    +91 98765 43210    ·    shopseed.com')
        _set_run(run, size=9.5, color=MUTED)

    return doc
